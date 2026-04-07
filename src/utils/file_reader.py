
from typing import Dict, List, Any, Optional, Union
from pathlib import Path

import os
import json
import yaml
import csv
import pandas as pd

from gen_ai_project_template.src.utils.fun_utils import seguimiento_funciones, msg_succ, msg_warn, msg_error, msg_info

@seguimiento_funciones
def read_file(
    file_path: str,
    encoding: str = 'utf-8',
    **kwargs
) -> Union[str, Dict, List, bytes, None]:
    """
    Leer cualquier archivo de forma genérica detectando automáticamente el formato.
    
    Args:
        file_path: Ruta al archivo
        encoding: Codificación del archivo (default: 'utf-8')
        **kwargs: Argumentos adicionales específicos del formato
    
    Returns:
        Contenido del archivo en el formato apropiado:
        - str: para TXT, XML
        - dict: para JSON, YAML
        - list: para CSV, JSONL, XLSX
        - bytes: para PDF, DOCX (contenido binario)
        - None: si hay error
    
    Example:
        >>> content = read_file("data.json")
        >>> content = read_file("data.csv", delimiter=';')
        >>> content = read_file("document.pdf")
    """
    try:
        if not os.path.exists(file_path):
            print(msg_error(f"Archivo no encontrado: {file_path}"))
            return None
        
        # Detectar extensión
        extension = Path(file_path).suffix.lower()
        
        # Mapeo de extensiones a funciones
        readers = {
            '.txt': read_txt_file,
            '.json': read_json_file,
            '.jsonl': read_jsonl_file,
            '.yaml': read_yaml_file,
            '.yml': read_yaml_file,
            '.csv': read_csv_file,
            '.tsv': lambda p, **kw: read_csv_file(p, delimiter='\t', **kw),
            '.xlsx': read_excel_file,
            '.xls': read_excel_file,
            '.pdf': read_pdf_file,
            '.docx': read_docx_file,
            '.xml': read_xml_file,
        }
        
        reader = readers.get(extension)
        
        if reader:
            # Binary formats (PDF, DOCX) don't accept encoding parameter
            if extension in ['.pdf', '.docx']:
                return reader(file_path, **kwargs)
            else:
                return reader(file_path, encoding=encoding, **kwargs)
        else:
            # Si no se reconoce la extensión, leer como texto plano
            print(msg_warn(f"Extensión no reconocida '{extension}', leyendo como texto plano"))
            return read_txt_file(file_path, encoding=encoding)
    
    except Exception as e:
        print(msg_error(f"Error al leer archivo: {e}"))
        return None


# ============================================================================
# FUNCIONES ESPECÍFICAS POR FORMATO
# ============================================================================

@seguimiento_funciones
def read_txt_file(
    file_path: str,
    encoding: str = 'utf-8'
) -> Optional[str]:
    """
    Leer archivo de texto plano.
    
    Args:
        file_path: Ruta al archivo
        encoding: Codificación del archivo
    
    Returns:
        Contenido del archivo como string
    
    Example:
        >>> content = read_txt_file("documento.txt")
    """
    try:
        with open(file_path, 'r', encoding=encoding) as f:
            content = f.read()
        print(msg_succ(f"Archivo TXT leído: {os.path.basename(file_path)}"))
        return content
    
    except Exception as e:
        print(msg_error(f"Error al leer archivo TXT: {e}"))
        return None


@seguimiento_funciones
def read_json_file(
    file_path: str,
    encoding: str = 'utf-8'
) -> Optional[Union[Dict, List]]:
    """
    Leer archivo JSON.
    
    Args:
        file_path: Ruta al archivo JSON
        encoding: Codificación del archivo
    
    Returns:
        Contenido del JSON como dict o list
    
    Example:
        >>> data = read_json_file("config.json")
    """
    try:
        with open(file_path, 'r', encoding=encoding) as f:
            data = json.load(f)
        print(msg_succ(f"Archivo JSON leído: {os.path.basename(file_path)}"))
        return data
    
    except json.JSONDecodeError as e:
        print(msg_error(f"Error al decodificar JSON: {e}"))
        return None
    except Exception as e:
        print(msg_error(f"Error al leer archivo JSON: {e}"))
        return None


@seguimiento_funciones
def read_jsonl_file(
    file_path: str,
    encoding: str = 'utf-8',
    skip_errors: bool = False
) -> List[Dict]:
    """
    Leer archivo JSONL (JSON Lines - una línea JSON por línea).
    
    Args:
        file_path: Ruta al archivo JSONL
        encoding: Codificación del archivo
        skip_errors: Si True, salta líneas con errores de JSON
    
    Returns:
        Lista de diccionarios, uno por cada línea
    
    Example:
        >>> data = read_jsonl_file("comentarios.jsonl")
    """
    try:
        data = []
        with open(file_path, 'r', encoding=encoding) as f:
            for line_num, line in enumerate(f, 1):
                line = line.strip()
                if not line:  # Saltar líneas vacías
                    continue
                
                try:
                    data.append(json.loads(line))
                except json.JSONDecodeError as e:
                    if skip_errors:
                        print(msg_warn(f"Error en línea {line_num}, saltando: {e}"))
                        continue
                    else:
                        print(msg_error(f"Error en línea {line_num}: {e}"))
                        return []
        
        print(msg_succ(f"Archivo JSONL leído: {len(data)} líneas desde {os.path.basename(file_path)}"))
        return data
    
    except Exception as e:
        print(msg_error(f"Error al leer archivo JSONL: {e}"))
        return []


@seguimiento_funciones
def read_yaml_file(
    file_path: str,
    encoding: str = 'utf-8'
) -> Optional[Dict]:
    """
    Leer archivo YAML.
    
    Args:
        file_path: Ruta al archivo YAML
        encoding: Codificación del archivo
    
    Returns:
        Contenido del YAML como dict
    
    Example:
        >>> config = read_yaml_file("config.yaml")
    """
    try:
        with open(file_path, 'r', encoding=encoding) as f:
            data = yaml.safe_load(f)
        print(msg_succ(f"Archivo YAML leído: {os.path.basename(file_path)}"))
        return data
    
    except yaml.YAMLError as e:
        print(msg_error(f"Error al decodificar YAML: {e}"))
        return None
    except Exception as e:
        print(msg_error(f"Error al leer archivo YAML: {e}"))
        return None


@seguimiento_funciones
def read_csv_file(
    file_path: str,
    encoding: str = 'utf-8',
    delimiter: str = ',',
    has_header: bool = True,
    as_dict: bool = True
) -> List[Union[Dict, List]]:
    """
    Leer archivo CSV.
    
    Args:
        file_path: Ruta al archivo CSV
        encoding: Codificación del archivo
        delimiter: Delimitador (default: ',')
        has_header: Si True, la primera fila es el encabezado
        as_dict: Si True, retorna lista de dicts. Si False, lista de listas
    
    Returns:
        Lista de diccionarios (si as_dict=True) o lista de listas
    
    Example:
        >>> data = read_csv_file("datos.csv")
        >>> data = read_csv_file("datos.tsv", delimiter='\\t')
    """
    try:
        data = []
        with open(file_path, 'r', encoding=encoding, newline='') as f:
            if as_dict and has_header:
                reader = csv.DictReader(f, delimiter=delimiter)
                data = list(reader)
            else:
                reader = csv.reader(f, delimiter=delimiter)
                data = list(reader)
        
        print(msg_succ(f"Archivo CSV leído: {len(data)} filas desde {os.path.basename(file_path)}"))
        return data
    
    except Exception as e:
        print(msg_error(f"Error al leer archivo CSV: {e}"))
        return []


@seguimiento_funciones
def read_excel_file(
    file_path: str,
    sheet_name: Union[str, int, None] = 0,
    header: Optional[int] = 0,
    as_dict: bool = True,
    **kwargs
) -> Union[List[Dict], List[List], Dict[str, List], None]:
    """
    Leer archivo Excel (XLSX, XLS).
    
    Args:
        file_path: Ruta al archivo Excel
        sheet_name: Nombre o índice de la hoja (0 = primera hoja, None = todas las hojas)
        header: Fila del encabezado (0 = primera fila, None = sin encabezado)
        as_dict: Si True, retorna lista de dicts
        **kwargs: Argumentos adicionales para pandas.read_excel (NO incluye encoding)
    
    Returns:
        - Si sheet_name es específico: Lista de dicts o listas
        - Si sheet_name es None: Dict con {nombre_hoja: datos}
    
    Example:
        >>> data = read_excel_file("datos.xlsx")
        >>> data = read_excel_file("datos.xlsx", sheet_name="Hoja2")
        >>> all_sheets = read_excel_file("datos.xlsx", sheet_name=None)
    """
    try:
        
        
        filtered_kwargs = {k: v for k, v in kwargs.items() if k != 'encoding'}
        df_or_dict = pd.read_excel(file_path, sheet_name=sheet_name, header=header, **filtered_kwargs)
        
        # Si se leyeron todas las hojas
        if isinstance(df_or_dict, dict):
            result = {}
            for sheet, df in df_or_dict.items():
                if as_dict:
                    result[sheet] = df.to_dict('records')
                else:
                    result[sheet] = df.values.tolist()
            print(msg_succ(f"Archivo Excel leído: {len(result)} hojas desde {os.path.basename(file_path)}"))
            return result
        
        # Si se leyó una sola hoja
        else:
            if as_dict:
                data = df_or_dict.to_dict('records')
            else:
                data = df_or_dict.values.tolist()
            print(msg_succ(f"Archivo Excel leído: {len(data)} filas desde {os.path.basename(file_path)}"))
            return data
    
    except ImportError:
        print(msg_error("pandas no está instalado. Instalar con: pip install pandas openpyxl"))
        return None
    except Exception as e:
        print(msg_error(f"Error al leer archivo Excel: {e}"))
        return None



@seguimiento_funciones
def read_pdf_file(
    file_path: str,
    extract_text: bool = True,
    page_numbers: Optional[List[int]] = None
) -> Union[str, bytes, None]:
    """
    Leer archivo PDF.
    
    Args:
        file_path: Ruta al archivo PDF
        extract_text: Si True, extrae texto. Si False, retorna bytes
        page_numbers: Lista de números de página a extraer (None = todas)
    
    Returns:
        Texto extraído (si extract_text=True) o bytes del archivo
    
    Example:
        >>> text = read_pdf_file("documento.pdf")
        >>> text = read_pdf_file("documento.pdf", page_numbers=[0, 1, 2])
    """
    try:
        if extract_text:
            import PyPDF2
            
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                total_pages = len(reader.pages)
                
                # Determinar qué páginas leer
                pages_to_read = page_numbers if page_numbers else range(total_pages)
                
                for page_num in pages_to_read:
                    if page_num < total_pages:
                        page = reader.pages[page_num]
                        text += page.extract_text() + "\n"
            
            print(msg_succ(f"PDF leído: {len(pages_to_read)} páginas desde {os.path.basename(file_path)}"))
            return text.strip()
        
        else:
            # Leer como bytes
            with open(file_path, 'rb') as f:
                content = f.read()
            print(msg_succ(f"PDF leído como bytes: {os.path.basename(file_path)}"))
            return content
    
    except ImportError:
        print(msg_error("PyPDF2 no está instalado. Instalar con: pip install PyPDF2"))
        return None
    except Exception as e:
        print(msg_error(f"Error al leer archivo PDF: {e}"))
        return None


@seguimiento_funciones
def read_docx_file(
    file_path: str,
    extract_text: bool = True
) -> Union[str, bytes, None]:
    """
    Leer archivo DOCX (Word).
    
    Args:
        file_path: Ruta al archivo DOCX
        extract_text: Si True, extrae texto. Si False, retorna bytes
    
    Returns:
        Texto extraído (si extract_text=True) o bytes del archivo
    
    Example:
        >>> text = read_docx_file("documento.docx")
    """
    try:
        if extract_text:
            from docx import Document
            
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            print(msg_succ(f"DOCX leído: {len(doc.paragraphs)} párrafos desde {os.path.basename(file_path)}"))
            return text.strip()
        
        else:
            # Leer como bytes
            with open(file_path, 'rb') as f:
                content = f.read()
            print(msg_succ(f"DOCX leído como bytes: {os.path.basename(file_path)}"))
            return content
    
    except ImportError:
        print(msg_error("python-docx no está instalado. Instalar con: pip install python-docx"))
        return None
    except Exception as e:
        print(msg_error(f"Error al leer archivo DOCX: {e}"))
        return None


@seguimiento_funciones
def read_xml_file(
    file_path: str,
    encoding: str = 'utf-8',
    as_dict: bool = False
) -> Union[str, Dict, None]:
    """
    Leer archivo XML.
    
    Args:
        file_path: Ruta al archivo XML
        encoding: Codificación del archivo
        as_dict: Si True, convierte XML a diccionario (requiere xmltodict)
    
    Returns:
        Contenido XML como string o dict
    
    Example:
        >>> content = read_xml_file("data.xml")
        >>> data = read_xml_file("data.xml", as_dict=True)
    """
    try:
        with open(file_path, 'r', encoding=encoding) as f:
            content = f.read()
        
        if as_dict:
            import xmltodict
            data = xmltodict.parse(content)
            print(msg_succ(f"XML leído como dict: {os.path.basename(file_path)}"))
            return data
        else:
            print(msg_succ(f"XML leído: {os.path.basename(file_path)}"))
            return content
    
    except ImportError:
        print(msg_error("xmltodict no está instalado. Instalar con: pip install xmltodict"))
        return None
    except Exception as e:
        print(msg_error(f"Error al leer archivo XML: {e}"))
        return None


# ============================================================================
# FUNCIONES UTILITARIAS
# ============================================================================

@seguimiento_funciones
def read_lines(
    file_path: str,
    encoding: str = 'utf-8',
    strip: bool = True,
    skip_empty: bool = False
) -> List[str]:
    """
    Leer archivo línea por línea.
    
    Args:
        file_path: Ruta al archivo
        encoding: Codificación del archivo
        strip: Si True, elimina espacios en blanco al inicio/final
        skip_empty: Si True, salta líneas vacías
    
    Returns:
        Lista de líneas
    
    Example:
        >>> lines = read_lines("archivo.txt")
        >>> lines = read_lines("archivo.txt", skip_empty=True)
    """
    try:
        with open(file_path, 'r', encoding=encoding) as f:
            lines = f.readlines()
        
        if strip:
            lines = [line.strip() for line in lines]
        
        if skip_empty:
            lines = [line for line in lines if line]
        
        print(msg_succ(f"Archivo leído: {len(lines)} líneas desde {os.path.basename(file_path)}"))
        return lines
    
    except Exception as e:
        print(msg_error(f"Error al leer líneas: {e}"))
        return []


@seguimiento_funciones
def get_file_info(file_path: str) -> Dict[str, Any]:
    """
    Obtener información sobre un archivo.
    
    Args:
        file_path: Ruta al archivo
    
    Returns:
        Diccionario con información del archivo
    
    Example:
        >>> info = get_file_info("documento.pdf")
        >>> print(f"Tamaño: {info['size_mb']} MB")
    """
    try:
        path = Path(file_path)
        
        if not path.exists():
            print(msg_error(f"Archivo no encontrado: {file_path}"))
            return {}
        
        stat = path.stat()
        
        info = {
            "name": path.name,
            "extension": path.suffix,
            "size_bytes": stat.st_size,
            "size_kb": round(stat.st_size / 1024, 2),
            "size_mb": round(stat.st_size / (1024 * 1024), 2),
            "absolute_path": str(path.absolute()),
            "parent_dir": str(path.parent),
            "exists": True,
            "is_file": path.is_file(),
            "is_dir": path.is_dir()
        }
        
        print(msg_info(f"Info obtenida: {path.name} ({info['size_kb']} KB)"))
        return info
    
    except Exception as e:
        print(msg_error(f"Error al obtener info del archivo: {e}"))
        return {}


@seguimiento_funciones
def read_multiple_files(
    file_paths: List[str],
    encoding: str = 'utf-8',
    **kwargs
) -> Dict[str, Any]:
    """
    Leer múltiples archivos de una vez.
    
    Args:
        file_paths: Lista de rutas a archivos
        encoding: Codificación de los archivos
        **kwargs: Argumentos adicionales para read_file
    
    Returns:
        Diccionario con {nombre_archivo: contenido}
    
    Example:
        >>> files = ["config.yaml", "data.json", "info.txt"]
        >>> contents = read_multiple_files(files)
    """
    results = {}
    
    for file_path in file_paths:
        file_name = os.path.basename(file_path)
        content = read_file(file_path, encoding=encoding, **kwargs)
        results[file_name] = content
    
    successful = sum(1 for v in results.values() if v is not None)
    print(msg_succ(f"Lectura múltiple completada: {successful}/{len(file_paths)} archivos"))
    
    return results


@seguimiento_funciones
def file_exists(file_path: str) -> bool:
    """
    Verificar si un archivo existe.
    
    Args:
        file_path: Ruta al archivo
    
    Returns:
        True si existe, False en caso contrario
    
    Example:
        >>> if file_exists("config.yaml"):
        ...     config = read_yaml_file("config.yaml")
    """
    exists = os.path.exists(file_path) and os.path.isfile(file_path)
    
    if exists:
        print(msg_info(f"Archivo existe: {os.path.basename(file_path)}"))
    else:
        print(msg_warn(f"Archivo no existe: {file_path}"))
    
    return exists